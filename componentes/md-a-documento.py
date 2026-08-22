#!/usr/bin/env python3
"""Convierte el subconjunto Markdown usado por Agente Smith a DOCX y PDF."""

from __future__ import annotations

import argparse
import html
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


@dataclass(frozen=True)
class Block:
    kind: str
    text: str
    level: int = 0


def parse_markdown(source: str) -> list[Block]:
    """Parse headings, paragraphs and ordered/unordered lists."""
    blocks: list[Block] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            blocks.append(Block("paragraph", " ".join(paragraph).strip()))
            paragraph.clear()

    for raw_line in source.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        ordered = re.match(r"^\d+[.)]\s+(.+)$", line)
        bullet = re.match(r"^[-*+]\s+(.+)$", line)
        if heading:
            flush()
            blocks.append(Block("heading", heading.group(2), len(heading.group(1))))
        elif ordered:
            flush()
            blocks.append(Block("ordered", ordered.group(1)))
        elif bullet:
            flush()
            blocks.append(Block("bullet", bullet.group(1)))
        else:
            paragraph.append(line)
    flush()
    return blocks


def inline_parts(text: str) -> Iterable[tuple[str, bool]]:
    """Yield text segments and whether each segment is bold."""
    position = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > position:
            yield text[position : match.start()], False
        yield match.group(1), True
        position = match.end()
    if position < len(text):
        yield text[position:], False


def _require(module: str, install_name: str):
    try:
        return __import__(module)
    except ImportError as exc:
        raise RuntimeError(
            f"Falta {install_name}. Instalá las dependencias con: "
            "python -m pip install -r componentes/requirements-documentos.txt"
        ) from exc


def write_docx(blocks: list[Block], destination: Path) -> None:
    docx = _require("docx", "python-docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = docx.Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    for block in blocks:
        if block.kind == "heading":
            paragraph = document.add_heading(level=min(block.level, 4))
        elif block.kind == "ordered":
            paragraph = document.add_paragraph(style="List Number")
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
        else:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for text, bold in inline_parts(block.text):
            run = paragraph.add_run(text)
            run.bold = bold

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def _find_unicode_font() -> Path:
    candidates = (
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/Library/Fonts/Arial Unicode.ttf"),
        Path.home() / "Library/Fonts/Arial.ttf",
    )
    for candidate in candidates:
        if candidate.exists():
            return candidate
    raise RuntimeError("No se encontró una fuente TrueType con soporte Unicode.")


def _pdf_markup(text: str) -> str:
    parts = []
    for value, bold in inline_parts(text):
        escaped = html.escape(value)
        parts.append(f"<b>{escaped}</b>" if bold else escaped)
    return "".join(parts)


def write_pdf(blocks: list[Block], destination: Path) -> None:
    _require("reportlab", "reportlab")
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    font_name = "AgenteSmithUnicode"
    pdfmetrics.registerFont(TTFont(font_name, str(_find_unicode_font())))
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "AgenteSmithBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=11,
        leading=16.5,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )
    headings = {
        level: ParagraphStyle(
            f"AgenteSmithHeading{level}",
            parent=styles[f"Heading{min(level, 4)}"],
            fontName=font_name,
        )
        for level in range(1, 7)
    }

    story = []
    list_items: list[ListItem] = []
    list_kind: str | None = None

    def flush_list() -> None:
        nonlocal list_items, list_kind
        if list_items:
            story.append(
                ListFlowable(
                    list_items,
                    bulletType="1" if list_kind == "ordered" else "bullet",
                    leftIndent=18,
                )
            )
            story.append(Spacer(1, 6))
        list_items = []
        list_kind = None

    for block in blocks:
        if block.kind in {"ordered", "bullet"}:
            if list_kind not in {None, block.kind}:
                flush_list()
            list_kind = block.kind
            list_items.append(ListItem(Paragraph(_pdf_markup(block.text), body)))
            continue
        flush_list()
        style = headings[block.level] if block.kind == "heading" else body
        story.append(Paragraph(_pdf_markup(block.text), style))
    flush_list()

    destination.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(destination),
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=destination.stem,
    )
    document.build(story)


def output_paths(source: Path, output: Path | None) -> tuple[Path, Path]:
    base = output if output is not None else source.with_suffix("")
    if base.suffix.lower() in {".docx", ".pdf"}:
        base = base.with_suffix("")
    return base.with_suffix(".docx"), base.with_suffix(".pdf")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path, help="Archivo Markdown de entrada")
    parser.add_argument("--output", "-o", type=Path, help="Ruta base de salida")
    args = parser.parse_args(argv)

    if not args.source.is_file():
        parser.error(f"No existe el archivo de entrada: {args.source}")
    blocks = parse_markdown(args.source.read_text(encoding="utf-8"))
    if not blocks:
        parser.error("El archivo Markdown no contiene texto para convertir")

    docx_path, pdf_path = output_paths(args.source, args.output)
    try:
        write_docx(blocks, docx_path)
        write_pdf(blocks, pdf_path)
    except RuntimeError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1
    print(f"Generados: {docx_path} | {pdf_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
